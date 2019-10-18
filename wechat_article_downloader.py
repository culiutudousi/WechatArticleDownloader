import datetime
import os
from enum import Enum
from re import findall, sub
from urllib.request import urlopen

from PIL import Image
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches
from docx.shared import Pt, Cm


class Singleton(object):
    _instance = None

    def __new__(cls, *args, **kw):
        if not cls._instance:
            cls._instance = super(Singleton, cls).__new__(cls, *args, **kw)
        return cls._instance


class Settings(Singleton):
    IMAGE_PATH = ''
    FONT = u'微软雅黑'
    FONT_SIZE = 10.5
    FIRST_LINE_INDENT = 0.25
    MAX_IMAGE_WIDTH = 6


class ParagraphType(Enum):
    SEPARATOR = 0
    TEXT = 1
    IMAGE = 2


class TextFormat:

    def __init__(self, bold=0, alignment_center=0, alignment_right=0, image_width=0):
        self.bold = bold
        self.alignment_center = alignment_center
        self.alignment_right = alignment_right
        self.image_width = image_width


class ArticleParagraph:

    def __init__(self, type, content='', text_format=TextFormat()):
        self.type = type
        self.content = content
        self.text_format = text_format


def parse_content(content, parent_text_format=TextFormat(bold=0,
                                                         alignment_center=0,
                                                         alignment_right=0)):
    text_format = TextFormat(bold=parent_text_format.bold,
                             alignment_center=parent_text_format.alignment_center,
                             alignment_right=parent_text_format.alignment_right)
    if not content:
        return []
    if isinstance(content, NavigableString):
        plain_text = content.replace('&nbsp;', '').strip()
        if plain_text:
            return [ArticleParagraph(ParagraphType.TEXT, plain_text, text_format)]
        return []
    if content.name == 'strong':
        text_format.bold = 1
    if 'style' in getattr(content, 'attrs', {}):
        style = content['style']
        if 'text-align: justify' in str(style):
            text_format.alignment_center = 0
            text_format.alignment_right = 0
        elif 'text-align: center' in str(style):
            text_format.alignment_center = 1
        elif 'text-align: right' in str(style):
            text_format.alignment_right = 1
    if content.name == 'img':
        if 'data-src' in getattr(content, 'attrs', {}):
            img_file = download_img(content['data-src'])
            if img_file:
                if 'style' in getattr(content, 'attrs', {}):
                    width = findall('width: (\d+?)px', content['style'])
                    if width and width[0]:
                        width_inch = float(width[0]) / 72
                        if width_inch < Settings().MAX_IMAGE_WIDTH:
                            text_format.image_width = width_inch
                return [ArticleParagraph(ParagraphType.IMAGE, img_file, text_format)]
        return [ArticleParagraph(ParagraphType.TEXT,
                                     "Can not download image from url: {}".format(
                                         content['data-src']))]
    parsed_content = []
    if content.children:
        for child in content.children:
            parsed_content += parse_content(child, TextFormat(bold=text_format.bold,
                                                              alignment_center=text_format.alignment_center,
                                                              alignment_right=text_format.alignment_right))
    if parsed_content and (content.name == 'section' or content.name == 'p'):
        if not parsed_content[-1].type == ParagraphType.SEPARATOR:
            parsed_content += [ArticleParagraph(ParagraphType.SEPARATOR)]
    return parsed_content


def get_img_format_from(url):
    lower_url = str(url).lower()
    img_formats = ['png', 'jpg', 'jpeg', 'gif', 'bmp']
    for img_format in img_formats:
        if lower_url.endswith(img_format):
            return img_format
    if lower_url.endswith('other'):
        return 'jpg'
    separator_index = max(url.find(separator, -6) for separator in '.=')
    if separator_index > 0:
        img_format = url[separator_index + 1:]
        print("Get new image format from image url: {}".format(img_format))
        return img_format
    print("Can not get image format from url: \n{}, try to use jpg".format(url))
    return 'jpg'


def download_img(url):
    img_format = get_img_format_from(url)
    if img_format:
        file_name = datetime.datetime.now().strftime('%Y%m%d_%H%M%S%f') + '.' + img_format
        file_path = os.path.join(Settings().IMAGE_PATH, file_name)
        with urlopen(url) as url_cache:
            with open(file_path, 'wb') as img_file:
                img_file.write(url_cache.read())
        reread_img = Image.open(file_path)
        if reread_img.mode != "RGB":
            reread_img = reread_img.convert('RGB')
        reread_img.save(file_path)
        return file_path
    print("Download image failed from url: \n{}".format(url))
    return None


def get_image_show_width(pix_width):
    max_pix_width = 600
    max_show_width = 6.0
    if pix_width > max_pix_width:
        return max_show_width
    return pix_width * max_show_width / max_pix_width


def validate_title(title):
    rstr = r"[\/\\\:\*\?\"\<\>\|\.]"  # '/ \ : * ? " < > | .'
    new_title = sub(rstr, "_", title)
    return new_title


def download_wechat_article_from(url):
    settings = Settings()

    with urlopen(url) as fp:
        html_content = fp.read().decode()

    soup = BeautifulSoup(html_content, "html5lib")
    soup.prettify()

    title_html = soup.find('h2', 'rich_media_title')
    title = title_html.string.strip()

    date_html = findall('s="(\d{4}-\d{2}-\d{2})"', str(soup))
    article_date = ''
    if date_html:
        article_date = date_html[0]

    img_path_base = 'images'
    img_path_article = os.path.join(img_path_base, ('wx ' + validate_title(title)))
    if not os.path.exists(img_path_base):
        os.makedirs(img_path_base)
    if not os.path.exists(img_path_article):
        os.makedirs(img_path_article)
    settings.IMAGE_PATH = img_path_article

    content_html = soup.find('div', 'rich_media_content')
    article_content = []
    for child in content_html.children:
        article_content += parse_content(child)

    return {'title': title,
            'date': article_date,
            'content': article_content}


def write_article_to_docx(article, url):
    settings = Settings()

    model_name = 'model.docx'
    if os.path.exists(model_name):
        document = Document(model_name)
    else:
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(1.91)
            section.right_margin = Cm(1.91)
    document.styles['Normal'].font.name = settings.FONT
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), settings.FONT)
    style = document.styles['Normal']
    style.font.size = Pt(settings.FONT_SIZE)

    current_paragraph = document.add_paragraph(url)
    current_paragraph.add_run(' (' + article['date'] + ')')
    paragraph_format = current_paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    current_paragraph = document.add_paragraph()
    current_paragraph.add_run(article['title']).bold = True
    current_paragraph_format = current_paragraph.paragraph_format
    current_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    current_paragraph = document.add_paragraph()
    for paragraph in article['content']:
        if paragraph.type == ParagraphType.TEXT:
            words = current_paragraph.add_run(paragraph.content)
            if paragraph.text_format.bold:
                words.bold = True
            if paragraph.text_format.alignment_center:
                current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif paragraph.text_format.alignment_right:
                current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                current_paragraph.paragraph_format.first_line_indent = Inches(settings.FIRST_LINE_INDENT)
        elif paragraph.type == ParagraphType.IMAGE:
            img_width = Image.open(paragraph.content).width
            if paragraph.text_format.image_width:
                show_width = paragraph.text_format.image_width
            else:
                show_width = get_image_show_width(img_width)
            # print('Adding img: ' + paragraph.content)
            document.add_picture(paragraph.content, width=Inches(show_width))
            if paragraph.text_format.alignment_center:
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif paragraph.type == ParagraphType.SEPARATOR:
            current_paragraph = document.add_paragraph()
    
    file_name = 'wx ' + validate_title(article['title']) + '.docx'
    document.save(file_name)
    return file_name


if __name__ == '__main__':

    debuging = False
    if debuging:
        print('===============================================')
        content = input('Please PASTE the link (url) of a wechat article, or Q to quit.\n')
        article = download_wechat_article_from(content)
        write_article_to_docx(article, url=content)
    else:
        while 1:
            print('===============================================\n')
            content = input('Please PASTE the link (url) of a wechat article, or Enter to quit.\n\n')
            if len(content) == 0:
                break
            try:
                print('===============================================\n')
                print('Please wait ...')
                article = download_wechat_article_from(url=content)
                file_name = write_article_to_docx(article, url=content)
                print('===============================================\n')
                print('Success! Saved as {}\n'.format(file_name))
            except:
                print('===================================================\n')
                print('Failed! Please press ENTER to TRY AGAIN.\n')
                input()
